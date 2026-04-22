from __future__ import annotations

import io
import os
import tempfile
from dataclasses import dataclass
from typing import List

import cv2
import numpy as np
from docx import Document
from fastapi import FastAPI, File, HTTPException, UploadFile
from fastapi.responses import FileResponse, HTMLResponse
from fastapi.staticfiles import StaticFiles
from fastapi.templating import Jinja2Templates
from paddleocr import PaddleOCR
from starlette.requests import Request


@dataclass
class OcrLine:
    text: str
    score: float
    x: int
    y: int
    w: int
    h: int


app = FastAPI(title="Handwriting OCR to Word")
app.mount("/static", StaticFiles(directory="static"), name="static")
templates = Jinja2Templates(directory="templates")

# PP-OCR supports Chinese handwriting reasonably well. Angle cls boosts robustness for photographed invoices.
ocr_engine = PaddleOCR(use_angle_cls=True, lang="ch", show_log=False)


def preprocess_image(image_bytes: bytes) -> np.ndarray:
    image_array = np.frombuffer(image_bytes, np.uint8)
    img = cv2.imdecode(image_array, cv2.IMREAD_COLOR)
    if img is None:
        raise HTTPException(status_code=400, detail="无法读取图片，请上传清晰的发票照片")

    # Denoise + contrast enhancement for faint handwriting.
    img = cv2.fastNlMeansDenoisingColored(img, None, 5, 5, 7, 21)
    lab = cv2.cvtColor(img, cv2.COLOR_BGR2LAB)
    l, a, b = cv2.split(lab)
    clahe = cv2.createCLAHE(clipLimit=2.5, tileGridSize=(8, 8))
    l = clahe.apply(l)
    img = cv2.cvtColor(cv2.merge((l, a, b)), cv2.COLOR_LAB2BGR)
    return img


def run_ocr(img: np.ndarray) -> List[OcrLine]:
    result = ocr_engine.ocr(img, cls=True)
    lines: List[OcrLine] = []

    if not result or not result[0]:
        return lines

    for item in result[0]:
        box = item[0]
        txt, score = item[1]
        xs = [int(p[0]) for p in box]
        ys = [int(p[1]) for p in box]
        x, y = min(xs), min(ys)
        w, h = max(xs) - x, max(ys) - y
        lines.append(OcrLine(text=txt.strip(), score=float(score), x=x, y=y, w=w, h=h))
    return lines


def looks_handwritten(line: OcrLine) -> bool:
    """
    Heuristic filter: printed text usually has very high confidence and regular short boxes.
    Handwriting tends to have lower confidence and variable box geometry.
    """
    if not line.text:
        return False
    if line.score < 0.92:
        return True
    if line.h > 34 and line.w / max(line.h, 1) > 1.8:
        return True
    return False


def sort_as_layout(lines: List[OcrLine]) -> List[OcrLine]:
    # Top-to-bottom then left-to-right reading order.
    return sorted(lines, key=lambda t: (t.y // 24, t.x))


def create_word(lines: List[OcrLine]) -> str:
    doc = Document()
    doc.add_heading("发票手写内容识别结果", level=1)

    if not lines:
        doc.add_paragraph("未识别到明确的中文手写内容。")
    else:
        for idx, line in enumerate(lines, 1):
            p = doc.add_paragraph()
            p.add_run(f"{idx}. {line.text}")

    fd, path = tempfile.mkstemp(prefix="handwriting_", suffix=".docx")
    os.close(fd)
    doc.save(path)
    return path


@app.get("/", response_class=HTMLResponse)
def home(request: Request):
    return templates.TemplateResponse("index.html", {"request": request})


@app.post("/api/extract")
async def extract_handwriting(file: UploadFile = File(...)):
    if not file.content_type or not file.content_type.startswith("image/"):
        raise HTTPException(status_code=400, detail="仅支持图片文件")

    content = await file.read()
    if len(content) > 15 * 1024 * 1024:
        raise HTTPException(status_code=400, detail="图片不能超过 15MB")

    img = preprocess_image(content)
    raw_lines = run_ocr(img)
    handwritten_lines = sort_as_layout([l for l in raw_lines if looks_handwritten(l)])

    doc_path = create_word(handwritten_lines)
    return FileResponse(
        doc_path,
        filename="手写识别结果.docx",
        media_type="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    )


if __name__ == "__main__":
    import uvicorn

    uvicorn.run("app:app", host="0.0.0.0", port=8000, reload=True)
